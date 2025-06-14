import os
import re
import sys
import traceback
from datetime import datetime
from docx import Document
from faster_whisper import WhisperModel
import tkinter as tk
from tkinter import filedialog

def add_timestamps_to_sentences(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()

    sentences = re.split(r'(?<!\b\w\.\w)(?<!\b\w\.)(?<=[.!?…])\s+', content)

    timestamped_content = ""
    start_time = 0.0
    duration = 2.0

    for idx, sentence in enumerate(sentences, start=1):
        if sentence.strip():
            end_time = start_time + duration
            start_timestamp = format_timestamp(start_time)
            end_timestamp = format_timestamp(end_time)
            timestamped_content += f"{idx}\n{start_timestamp} --> {end_timestamp}\n{sentence.strip()}\n\n"
            start_time = end_time

    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(timestamped_content)

    print(f"Timestamps added to {output_file}")

def convert_srt_to_docx(srt_file_path):
    """
    Converts an .srt subtitle file into a .docx file with format:
    (timestamp \n sentence)
    """
    if not os.path.exists(srt_file_path):
        print(f"SRT file not found: {srt_file_path}")
        return

    # Create a new Word document
    doc = Document()

    with open(srt_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split the content into subtitle blocks
    blocks = re.split(r'\n\s*\n', content.strip())

    for block in blocks:
        lines = block.strip().split('\n')
        if len(lines) >= 2:
            # The second line usually contains the timestamp
            time_line = lines[1] if '-->' in lines[1] else None
            # All lines after the timestamp line are subtitle text
            text_lines = lines[2:] if time_line else lines[1:]
            text = " ".join(text_lines).strip()

            if time_line and text:
                doc.add_paragraph(time_line)
                doc.add_paragraph(text)
                doc.add_paragraph("")  # Add an empty line between entries

    # Save the Word document next to the original .srt file
    output_path = os.path.splitext(srt_file_path)[0] + ".docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")

def generate_srt_from_audio_word_level(audio_path, language="ru", model=None):
    if model is None:
        raise ValueError("WhisperModel instance must be provided")
    segments, _ = model.transcribe(
        audio_path,
        language=language,
        word_timestamps=True,
        beam_size=5,
        # batch_size=8,
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=500))

    words = []
    for segment in segments:
        if segment.words:
            for word_info in segment.words:
                if word_info.start is not None and word_info.end is not None:
                    words.append({
                        'word': word_info.word,
                        'start': word_info.start,
                        'end': word_info.end
                    })

    sentences = group_words_into_sentences(words)
    srt_path = os.path.splitext(audio_path)[0] + "_wordlevel.srt"

    with open(srt_path, "w", encoding="utf-8") as f:
        for i, (start, end, text) in enumerate(sentences, start=1):
            f.write(f"{i}\n")
            f.write(f"{format_timestamp(start)} --> {format_timestamp(end)}\n")
            f.write(f"{text}\n\n")

    print(f"SRT saved to {srt_path}")
    return srt_path

def format_timestamp(seconds):
    ms = int((seconds % 1) * 1000)
    s = int(seconds)
    h = s // 3600
    m = (s % 3600) // 60
    s = s % 60
    return f"{h:02}:{m:02}:{s:02},{ms:03}"

def group_words_into_sentences(words, max_pause=1.0):
    sentences = []
    current_sentence = []
    start_time = None

    for i, word_info in enumerate(words):
        word = word_info['word']
        start = word_info['start']
        end = word_info['end']

        if start_time is None:
            start_time = start

        current_sentence.append((word, start, end))

        next_word_start = words[i + 1]['start'] if i + 1 < len(words) else None
        pause = next_word_start - end if next_word_start is not None else 0

        is_sentence_end = re.match(r'.*[.!?…]$', word.strip()) or pause > max_pause or i == len(words) - 1

        if is_sentence_end:
            sentence_text = " ".join(w[0] for w in current_sentence).strip()
            sentence_start = current_sentence[0][1]
            sentence_end = current_sentence[-1][2]
            sentences.append((sentence_start, sentence_end, sentence_text))
            current_sentence = []
            start_time = None

    return sentences

def log_error(error_message):
    log_file = "error_log.txt"
    with open(log_file, "a", encoding="utf-8") as file:
        file.write(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {error_message}\n")
    print(f"An error occurred. Details have been logged to {log_file}.")

def main():
    try:
        print("Author: Klimentsi Katsko (@leopalladium)")
        print("Welcome to the Time Stamper script!")

        # Ask user to select root directory
        root = tk.Tk()
        root.withdraw()
        print("Please select the root directory containing audio files...")
        root_dir = filedialog.askdirectory()
        if not root_dir:
            print("No directory selected. Exiting.")
            return
        # TODO: Make the option to convert srt to docx
        # Recursively find audio files
        audio_files = []
        for dirpath, _, files in os.walk(root_dir):
            for f in files:
                if f.lower().endswith(('.mp3', '.wav', '.m4a')):
                    audio_files.append(os.path.join(dirpath, f))

        if not audio_files:
            print("No audio files found in the selected directory.")
            return
        # TODO: Ask user if they want to process all files or select specific ones
        # TODO: Ask user if they want to generate docx
        print("Select the Whisper model to use:")
        print("1: tiny\n2: base\n3: small\n4: medium\n5: large")
        model_choice = input("Enter your choice (1-5): ").strip()
        model_name = {"1": "tiny", "2": "base", "3": "small", "4": "medium", "5": "large"}.get(model_choice, "base")

        language = input("Enter the language code for transcription (e.g., 'en', 'ru'): ").strip()
        model = WhisperModel(model_name, device="cuda", compute_type="int8_float16") # Adjust compute_type as needed

        for idx, file in enumerate(audio_files, start=1):
            print(f"{idx}: {file}")

        selected_files_input = input("Enter the numbers of the audio files to process (comma-separated), or type 'all' to process all: ").strip()
        if selected_files_input.lower() == "all":
            selected_files = audio_files
        else:
            selected_indices = [int(i.strip()) - 1 for i in selected_files_input.split(',')]
            selected_files = [audio_files[i] for i in selected_indices if 0 <= i < len(audio_files)]

        for file in selected_files:
            # Output SRT in the same folder as audio
            srt_path = os.path.splitext(file)[0] + "_wordlevel.srt"
            txt_file = generate_srt_from_audio_word_level(file, language, model)
            # Move SRT to correct folder if needed
            if os.path.abspath(txt_file) != os.path.abspath(srt_path):
                os.rename(txt_file, srt_path)
            print(f"Generated SRT file: {srt_path}")

            # print(f"Generate .docx file for {file}? (yes/no): ", end='', flush=True)
            # answer = input().strip().lower()
            # if answer == "yes":
            #     convert_srt_to_docx(srt_path)

    except Exception:
        error_message = traceback.format_exc()
        log_error(error_message)
        print("An unexpected error occurred. Please check the error log for details.")

    input("Task completed. Press Enter to exit...")

if __name__ == "__main__":
    main()
